"""Generate realistic pharmaceutical complaint documents for demonstration.

    python scripts/generate_samples.py

Writes PDF / EML / TXT / DOCX files into ``backend/samples/``. Each one is designed
to exercise a different part of the agent:

  01  contamination email        -> Critical risk, regulatory reportable
  02  packaging defect PDF       -> High risk, supplier-quality root causes
  03  labelling error text       -> line-clearance CAPA path
  04  adverse event PDF          -> pharmacovigilance escalation
  05  duplicate of 01            -> duplicate detection (same batch, same defect)
  06  sparse complaint           -> completeness checker flags missing mandatories
"""
from __future__ import annotations

from email.message import EmailMessage
from pathlib import Path

SAMPLES = Path(__file__).resolve().parent.parent / "samples"


# --------------------------------------------------------------------------
# Content
# --------------------------------------------------------------------------

CONTAMINATION_EMAIL = """Dear Quality Assurance Team,

I am writing to formally raise a serious product quality complaint on behalf of St Jude Hospital Pharmacy.

Customer Name: St Jude Hospital Pharmacy
Contact: p.marsden@stjudehospital.org / +44 161 555 0142
Product Name: Amoxicillin Trihydrate Capsules
Strength: 500 mg
Batch No: AMX-2451-B
Mfg Date: 05/01/2025
Expiry Date: 04/01/2027
Quantity Affected: 240 units
Complaint Date: 12/03/2026

Description: During routine dispensing on the morning of 11 March our senior pharmacist
identified black particulate matter clearly visible inside twelve capsules drawn from this
batch. The particles are approximately 0.5-1 mm and appear embedded in the powder fill
rather than sitting loose in the capsule shell. Two patients on the respiratory ward
received doses before the issue was identified; one subsequently reported nausea and was
kept under observation for four hours as a precaution. No hospitalisation was required.

We have quarantined the remaining 228 units and sequestered the affected blister strips
for your investigation. Photographs are available on request. Given the potential patient
safety implications we request an urgent response and confirmation of whether other
batches from the same campaign are affected.

Regards,
Dr Priya Marsden
Chief Pharmacist, St Jude Hospital
"""

DUPLICATE_TEXT = """PHARMACOVIGILANCE / QUALITY COMPLAINT INTAKE FORM

Complaint Source: Distributor
Customer Name: Northgate Medical Distributors Ltd
Contact: qa@northgatemed.co.uk
Product Name: Amoxicillin Trihydrate Capsules
Strength: 500 mg
Batch No: AMX-2451-B
Mfg Date: 05/01/2025
Expiry Date: 04/01/2027
Quantity Affected: 60 units
Complaint Date: 16/03/2026

Description: A downstream pharmacy customer returned three sealed packs after reporting
dark foreign particles visible within the capsule contents. On inspection at our depot we
confirmed the presence of black particulate matter in at least five capsules across two
packs. All units are from batch AMX-2451-B. We understand a similar issue may already have
been reported by a hospital customer earlier this month. Remaining depot stock of this
batch has been placed on hold pending your instruction.

Reported by: Alan Whitcombe, QA Manager, Northgate Medical Distributors Ltd
"""

PACKAGING_PDF_LINES = [
    ("h1", "CUSTOMER COMPLAINT NOTIFICATION"),
    ("h2", "Meridian Pharmaceuticals - Quality Assurance Department"),
    ("sp", ""),
    ("kv", "Complaint Source: Customer Portal"),
    ("kv", "Customer Name: Riverside Community Pharmacy Group"),
    ("kv", "Contact: complaints@riversidepharmacy.com"),
    ("kv", "Complaint Date: 04/02/2026"),
    ("sp", ""),
    ("h3", "PRODUCT DETAILS"),
    ("kv", "Product Name: Metformin Hydrochloride Prolonged-Release Tablets"),
    ("kv", "Strength: 1000 mg"),
    ("kv", "Batch No: MET-7712-A"),
    ("kv", "Mfg Date: 18/09/2025"),
    ("kv", "Expiry Date: 17/09/2027"),
    ("kv", "Quantity Affected: 96 units"),
    ("sp", ""),
    ("h3", "NATURE OF COMPLAINT"),
    ("p", "Description: Eight blister strips within the received consignment show incomplete "
          "heat sealing along the lower edge. The foil lidding lifts away from the base web "
          "under light finger pressure, leaving several tablet pockets effectively open to "
          "the atmosphere. Two of the exposed tablets show surface discolouration and a "
          "chalky texture consistent with moisture ingress."),
    ("p", "The defect appears confined to strips bearing the same overprint time code, which "
          "suggests an intermittent sealing station issue rather than a whole-batch failure. "
          "We have withdrawn the affected packs from the dispensing shelf and request a "
          "formal investigation with a written response within ten working days."),
    ("sp", ""),
    ("kv", "Reported by: Helen Okafor, Superintendent Pharmacist"),
]

ADVERSE_EVENT_PDF_LINES = [
    ("h1", "FIELD ALERT - ADVERSE EVENT COMPLAINT"),
    ("h2", "Received via Sales Representative - Priority Handling Requested"),
    ("sp", ""),
    ("kv", "Complaint Source: Sales Representative"),
    ("kv", "Customer Name: Lakeside General Hospital"),
    ("kv", "Contact: dr.eze@lakesidegen.health"),
    ("kv", "Complaint Date: 22/05/2026"),
    ("sp", ""),
    ("h3", "PRODUCT DETAILS"),
    ("kv", "Product Name: Ceftriaxone Sodium for Injection"),
    ("kv", "Strength: 1 g vial"),
    ("kv", "Batch No: CFX-3390-D"),
    ("kv", "Mfg Date: 02/11/2025"),
    ("kv", "Expiry Date: 01/11/2027"),
    ("kv", "Quantity Affected: 18 vials"),
    ("sp", ""),
    ("h3", "EVENT NARRATIVE"),
    ("p", "Description: Three patients on the surgical ward developed marked injection site "
          "inflammation and pyrexia within two hours of administration from vials of this "
          "batch. One patient, a 68-year-old male, deteriorated and was hospitalised in the "
          "high dependency unit for 48 hours with a suspected pyrogenic reaction. The "
          "attending consultant considers a product-related endotoxin issue plausible and "
          "has escalated internally."),
    ("p", "On reconstitution, staff noted that the solution took noticeably longer to clear "
          "than usual and had a faint yellow tint not seen with previous batches. All "
          "remaining vials from this batch have been quarantined and retained under "
          "refrigeration for collection."),
    ("sp", ""),
    ("kv", "Reported by: Marcus Bell, Territory Manager (on behalf of Dr A. Eze)"),
]

LABELLING_TEXT = """From: warehouse@carltonwholesale.co.uk
To: complaints@meridianpharma.com
Subject: Labelling discrepancy - Levothyroxine 50mcg - batch LEV-8823-C
Date: 09 April 2026

Hello,

Raising a labelling issue found during goods-in inspection this morning.

Customer Name: Carlton Wholesale Distribution
Contact: warehouse@carltonwholesale.co.uk
Product Name: Levothyroxine Sodium Tablets
Strength: 50 mcg
Batch No: LEV-8823-C
Mfg Date: 12/12/2025
Expiry Date: 11/12/2027
Quantity Affected: 1440 units
Complaint Date: 09/04/2026

Description: Approximately 60 cartons in this delivery carry an outer label showing a
strength of 100 mcg while the blister foil inside and the carton end-flap overprint both
correctly state 50 mcg. The barcode on the mislabelled cartons scans to the 100 mcg
product code in our warehouse system, which means these units would have been picked and
shipped as the wrong strength had the discrepancy not been caught at inspection.

This is a wrong-strength labelling error with clear potential for patient harm. The entire
delivery has been placed on hold and segregated. We need written confirmation of the scope
of the affected batch before we can release any stock.

Regards,
Sandra Iqbal
Goods-In Supervisor, Carlton Wholesale Distribution
"""

SPARSE_TEXT = """Phone call log - 30/06/2026

Caller rang the customer care line about tablets that "look wrong". Said the colour was
patchy on some of them and a couple were chipped. Did not have the pack to hand so could
not give a batch number or expiry. Said the product was a blood pressure tablet but was
not sure of the strength. Advised to call back with the carton details.

Call handled by: R. Ncube, Customer Care
"""


# --------------------------------------------------------------------------
# Writers
# --------------------------------------------------------------------------


def write_eml() -> Path:
    message = EmailMessage()
    message["From"] = "Dr Priya Marsden <p.marsden@stjudehospital.org>"
    message["To"] = "complaints@meridianpharma.com"
    message["Cc"] = "qa.director@stjudehospital.org"
    message["Subject"] = "URGENT: Foreign particulate matter - Amoxicillin 500mg batch AMX-2451-B"
    message["Date"] = "Thu, 12 Mar 2026 09:14:22 +0000"
    message.set_content(CONTAMINATION_EMAIL)

    path = SAMPLES / "01_contamination_complaint.eml"
    path.write_bytes(message.as_bytes())
    return path


def write_pdf(filename: str, lines: list[tuple[str, str]]) -> Path:
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    base = getSampleStyleSheet()
    styles = {
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontSize=15, spaceAfter=4),
        "h2": ParagraphStyle("h2", parent=base["Normal"], fontSize=10, textColor="#555555", spaceAfter=10),
        "h3": ParagraphStyle("h3", parent=base["Heading3"], fontSize=11, spaceBefore=8, spaceAfter=4),
        "kv": ParagraphStyle("kv", parent=base["Normal"], fontSize=10, leading=15),
        "p": ParagraphStyle("p", parent=base["Normal"], fontSize=10, leading=15, alignment=TA_JUSTIFY, spaceAfter=8),
    }

    path = SAMPLES / filename
    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm, topMargin=20 * mm, bottomMargin=20 * mm,
        title=filename,
    )

    flowables = []
    for kind, text in lines:
        if kind == "sp":
            flowables.append(Spacer(1, 6))
        else:
            flowables.append(Paragraph(text, styles[kind]))
    doc.build(flowables)
    return path


def write_docx() -> Path:
    import docx

    document = docx.Document()
    document.add_heading("Customer Complaint Record", level=1)
    document.add_paragraph("Meridian Pharmaceuticals - Quality Assurance")

    table = document.add_table(rows=0, cols=2)
    for key, value in [
        ("Complaint Source", "Regulatory Body"),
        ("Customer Name", "State Drug Control Administration"),
        ("Contact", "inspection.cell@sdca.gov.in"),
        ("Product Name", "Paracetamol Oral Suspension"),
        ("Strength", "250 mg/5 ml"),
        ("Batch No", "PCM-1180-K"),
        ("Mfg Date", "14/07/2025"),
        ("Expiry Date", "13/07/2027"),
        ("Quantity Affected", "300 units"),
        ("Complaint Date", "18/06/2026"),
    ]:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    document.add_heading("Description", level=2)
    document.add_paragraph(
        "During a routine market sample inspection, the drawn sample of this batch failed "
        "the assay test, returning 88.4% of the labelled paracetamol content against a "
        "specification of 95.0-105.0%. The sample also showed a marginally elevated "
        "4-aminophenol degradation impurity at 0.12% against a limit of 0.10%. The result "
        "is an out-of-specification finding on a marketed batch and the licence holder is "
        "directed to investigate, respond within 14 days, and confirm whether a recall of "
        "the batch is warranted."
    )
    document.add_paragraph("Issued by: Drug Inspector, Zone 3, State Drug Control Administration")

    path = SAMPLES / "07_regulatory_oos_notice.docx"
    document.save(str(path))
    return path


def main() -> None:
    SAMPLES.mkdir(parents=True, exist_ok=True)

    written = [
        write_eml(),
        write_pdf("02_packaging_defect_complaint.pdf", PACKAGING_PDF_LINES),
        write_pdf("04_adverse_event_field_alert.pdf", ADVERSE_EVENT_PDF_LINES),
        write_docx(),
    ]

    for name, content in [
        ("03_labelling_error_complaint.txt", LABELLING_TEXT),
        ("05_duplicate_of_contamination.txt", DUPLICATE_TEXT),
        ("06_incomplete_phone_complaint.txt", SPARSE_TEXT),
    ]:
        path = SAMPLES / name
        path.write_text(content, encoding="utf-8")
        written.append(path)

    print(f"Wrote {len(written)} sample documents to {SAMPLES}")
    for path in sorted(written):
        print(f"  {path.name:<42} {path.stat().st_size:>7,} bytes")


if __name__ == "__main__":
    main()
