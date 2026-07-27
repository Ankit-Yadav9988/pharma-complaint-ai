"""Extract plain text from uploaded complaint documents.

Supports PDF, DOCX, TXT/MD and EML. This is intentionally lightweight — no OCR —
which matches the brief: the goal is to feed the agent readable text, not to build
a production document pipeline. Scanned/image-only PDFs are reported clearly rather
than silently returning nothing.
"""
from __future__ import annotations

import io
from email import policy
from email.parser import BytesParser
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".eml", ".msg"}


class UnsupportedDocument(Exception):
    pass


class EmptyDocument(Exception):
    pass


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — a single bad page shouldn't kill the upload
            continue
    return "\n\n".join(p for p in pages if p.strip())


def _parse_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(": ".join(cells))
    return "\n".join(parts)


def _parse_eml(data: bytes) -> str:
    message = BytesParser(policy=policy.default).parsebytes(data)

    header_lines = []
    for header in ("From", "To", "Cc", "Subject", "Date"):
        value = message.get(header)
        if value:
            header_lines.append(f"{header}: {value}")

    body = ""
    if message.is_multipart():
        for part in message.walk():
            if part.get_content_type() == "text/plain" and "attachment" not in str(
                part.get("Content-Disposition", "")
            ):
                body += part.get_content()
    else:
        body = message.get_content()

    if not body.strip():
        # Fall back to the HTML part with tags stripped.
        import re

        for part in message.walk():
            if part.get_content_type() == "text/html":
                body = re.sub(r"<[^>]+>", " ", part.get_content())
                break

    return "\n".join(header_lines) + "\n\n" + str(body).strip()


def parse_document(filename: str, data: bytes) -> str:
    """Return the plain text of *data*, dispatching on the file extension."""
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocument(
            f"'{suffix or filename}' is not supported. Upload a PDF, DOCX, TXT or EML file."
        )

    if suffix == ".pdf":
        text = _parse_pdf(data)
        if not text.strip():
            raise EmptyDocument(
                "No selectable text found in this PDF. It is most likely a scan — "
                "paste the complaint text instead, or upload a text-based PDF."
            )
    elif suffix == ".docx":
        text = _parse_docx(data)
    elif suffix in {".eml", ".msg"}:
        text = _parse_eml(data)
    else:
        text = data.decode("utf-8", errors="replace")

    if not text.strip():
        raise EmptyDocument("The document appears to be empty.")

    return text.strip()
